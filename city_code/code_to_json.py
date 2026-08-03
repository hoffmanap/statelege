"""
City code converter — Word/Excel -> per-section JSON
El Paso Titles 18 (Building), 19 (Subdivision), 20 (Zoning), 21 (SmartCode)

Two input shapes are handled differently:

  1. NARRATIVE SECTIONS (.docx) — ordinance text like 20.10.625 or 20.10.035.
     Split on section-number headers (e.g. "20.10.625" or "20.10.363") and
     write one JSON file per section.

  2. TABLES (.docx with an embedded table, or .xlsx) — e.g. Appendix B's
     density/dimensional standards table. These don't split into narrative
     sections; instead each row (a zoning-district/use combination) becomes
     one JSON record, since that's the actual unit of comparison a bill
     conflict-check needs (e.g. "does this bill's setback rule conflict with
     R-3 single-family setbacks specifically").

Usage:
    pip install python-docx openpyxl --break-system-packages
    python code_to_json.py --input /path/to/docx_or_xlsx_files --output ./city_code

Each narrative section file looks like:
    {
      "section": "20.10.625",
      "title": "Small lot development",
      "title_num": 20,
      "source_doc": "20_10_625___Small_lot_development_.docx",
      "text": "A. The standards in this section supersede..."
    }

Each table row file looks like:
    {
      "table": "Appendix B",
      "title_num": 20,
      "zoning_district": "R-3",
      "permitted_use": "Single-family dwelling",
      "standards": {
          "Minimum Lot Area (square feet)": "6,000",
          "Minimum Average Lot Width (in feet)": "60",
          ...
      },
      "source_doc": "Appendix_B___TABLE_OF_DENSITY_AND_DIMENSIONAL_STANDARDS.docx"
    }
"""

import argparse
import json
import re
from pathlib import Path

from docx import Document

# Matches section numbers like 20.10.625 or 18.04.100 at the start of a line,
# optionally followed by a title. Adjust if a title uses a different pattern
# (e.g. Title 19 may number sections differently — check before a full run).
SECTION_HEADER_RE = re.compile(r"^(\d{2}\.\d{2}\.\d{2,4})\s*[-–—]?\s*(.*)$")


def title_num_from_section(section_number):
    """20.10.625 -> 20"""
    return int(section_number.split(".")[0])


def read_docx_paragraphs(path):
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)
    return lines


def has_table(path):
    doc = Document(path)
    return len(doc.tables) > 0


def split_narrative_sections(lines, source_doc):
    """
    Walk paragraph lines, start a new section whenever a line matches the
    section-header pattern, and accumulate everything after it as that
    section's body text until the next header.
    """
    sections = []
    current = None

    for line in lines:
        m = SECTION_HEADER_RE.match(line)
        if m:
            if current:
                sections.append(current)
            section_number = m.group(1)
            title = m.group(2).strip().rstrip(".")
            current = {
                "section": section_number,
                "title": title,
                "title_num": title_num_from_section(section_number),
                "source_doc": source_doc,
                "text_lines": [line],
            }
        elif current:
            current["text_lines"].append(line)
        # Lines before the first recognized header (boilerplate, doc titles)
        # are dropped — they're not part of any addressable code section.

    if current:
        sections.append(current)

    for s in sections:
        s["text"] = "\n".join(s.pop("text_lines"))

    return sections


def convert_docx_table(path, output_dir, table_label="Appendix B", title_num=20):
    """
    Extract a docx table (e.g. Appendix B) into one JSON record per row.
    Assumes the first table row is the header row with column names, and
    the first two columns identify zoning district and permitted use —
    matches the structure of Appendix B as provided. Adjust column indices
    if a different table has a different layout.

    Rows where the 'permitted use' or 'zoning district' cell is much
    longer than a normal label (e.g. a full legal description or overlay
    boundary paragraph embedded in an unrelated table in the same
    document) are skipped rather than converted — those aren't standards
    rows, and forcing them through produces garbage records and, on
    Windows, filenames that blow past the path length limit.
    """
    MAX_LABEL_LEN = 80

    doc = Document(path)
    written = 0
    skipped = 0

    for table in doc.tables:
        rows = table.rows
        if len(rows) < 2:
            continue

        headers = [cell.text.strip() for cell in rows[0].cells]

        for row in rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if not any(cells):
                continue

            # Column layout observed in Appendix B: row-label, district, min
            # district area, permitted use, then standards columns. Adjust
            # if a different table (e.g. Title 21 SmartCode) is shaped
            # differently.
            zoning_district = cells[1] if len(cells) > 1 else ""
            permitted_use = cells[3] if len(cells) > 3 else ""

            if len(permitted_use) > MAX_LABEL_LEN or len(zoning_district) > 40:
                skipped += 1
                continue

            standards = {}
            for h, c in zip(headers, cells):
                if h and h not in ("A", "B", "C"):  # skip bare row-label columns
                    standards[h] = c

            record = {
                "table": table_label,
                "title_num": title_num,
                "zoning_district": zoning_district,
                "permitted_use": permitted_use,
                "standards": standards,
                "source_doc": path.name,
            }

            safe_district = re.sub(r"[^\w\-]", "_", zoning_district or "row")[:25]
            safe_use = re.sub(r"[^\w\-]", "_", (permitted_use or "use"))[:25]
            out_path = output_dir / f"{table_label.replace(' ', '_')[:20]}_{safe_district}_{safe_use}_{written}.json"
            out_path.write_text(json.dumps(record, indent=2))
            written += 1

    return written, skipped


def run(input_dir, output_root):
    input_dir = Path(input_dir)
    output_root = Path(output_root)
    output_root.mkdir(exist_ok=True)

    docx_files = sorted(input_dir.glob("*.docx"))
    if not docx_files:
        print(f"No .docx files found in {input_dir}")
        return

    total_sections = 0
    total_rows = 0

    for path in docx_files:
        print(f"Processing {path.name}")

        # Narrative sections: extract from paragraph text regardless of
        # whether the file ALSO contains a table. A single docx (like a
        # full Title 20 export) can have both ordinance text and an
        # embedded table — these aren't mutually exclusive.
        lines = read_docx_paragraphs(path)
        sections = split_narrative_sections(lines, path.name)
        if sections:
            title_num = sections[0]["title_num"]
            title_out = output_root / f"title_{title_num}"
            title_out.mkdir(exist_ok=True)
            for s in sections:
                out_path = title_out / f"{s['section']}.json"
                out_path.write_text(json.dumps(s, indent=2))
            total_sections += len(sections)
            print(f"  wrote {len(sections)} narrative section(s) to {title_out}")
        else:
            print(f"  no narrative section headers found in {path.name}")

        # Tables: extract separately if present. Skips rows whose
        # "permitted use" cell is implausibly long (a sign it's not a
        # standards table row but something like a legal description or
        # overlay boundary text) rather than trying to force it into the
        # Appendix-B-shaped record.
        if has_table(path):
            table_out = output_root / "tables"
            table_out.mkdir(exist_ok=True)
            n, skipped = convert_docx_table(path, table_out)
            total_rows += n
            print(f"  wrote {n} table row(s) to {table_out}"
                  + (f" ({skipped} row(s) skipped — didn't match expected table shape)" if skipped else ""))

    print(f"\nDone. {total_sections} narrative section(s), {total_rows} table row(s) written.")
    print("Spot-check a few JSON files before committing — section-header "
          "regex and table column indices are pattern-matched, not guaranteed "
          "correct for every document layout.")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert city code docx/xlsx into per-section JSON")
    parser.add_argument("--input", required=True, help="Directory containing source .docx files")
    parser.add_argument("--output", default="./city_code", help="Output root directory (city_code/)")
    args = parser.parse_args()
    run(args.input, args.output)